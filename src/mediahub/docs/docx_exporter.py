from pathlib import Path

from src.mediahub.docs.markdown_tools import MarkdownTools


class DocxExporter:
    def __init__(self):
        self.tools = MarkdownTools()

    def write(self, book, docs, figures, lookup, target: Path):
        try:
            from docx import Document
            from docx.shared import Inches
        except Exception:
            print("⚠ python-docx fehlt. DOCX wird übersprungen.")
            return

        document = Document()
        document.add_heading(book.title, 0)

        document.add_heading("Inhaltsverzeichnis", level=1)
        for doc in docs:
            document.add_paragraph(doc["title"], style="List Bullet")

        if figures:
            document.add_heading("Abbildungsverzeichnis", level=1)
            for fig in figures:
                document.add_paragraph(
                    f"Abbildung {fig['number']}: {fig['caption']}",
                    style="List Bullet",
                )

        for doc in docs:
            for line in doc["text"].splitlines():
                clean = line.strip()
                ref = self.tools.image_reference(clean)

                if ref:
                    image_name, _caption = ref
                    fig = self.tools.image_figure(doc, lookup, image_name)

                    if fig and fig["exists"]:
                        try:
                            document.add_picture(str(fig["path"]), width=Inches(6.2))
                            document.add_paragraph(
                                f"Abbildung {fig['number']}: {fig['caption']}"
                            )
                        except Exception as error:
                            document.add_paragraph(
                                f"Bild konnte nicht eingefügt werden: "
                                f"{fig['path'].name} ({error})"
                            )
                    elif fig:
                        document.add_paragraph(
                            f"Abbildung {fig['number']} fehlt: {fig['path'].name}"
                        )

                    continue

                if clean.startswith("# "):
                    document.add_heading(clean[2:], level=1)
                elif clean.startswith("## "):
                    document.add_heading(clean[3:], level=2)
                elif clean.startswith("### "):
                    document.add_heading(clean[4:], level=3)
                elif clean.startswith("- "):
                    document.add_paragraph(clean[2:], style="List Bullet")
                elif clean:
                    document.add_paragraph(clean)

        document.save(target)
